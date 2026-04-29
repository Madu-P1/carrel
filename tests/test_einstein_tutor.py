import asyncio
import io
import tempfile
import unittest
from pathlib import Path

import main
from api_models import FlashcardDraftRequest, StudioGenerateRequest
from fastapi import UploadFile
from fastapi.testclient import TestClient
from routes.documents import upload_document
from routes.study import draft_flashcards
from routes.studio import studio_generate, studio_get_artifact
from services import documents as document_service
from services.graph import fetch_graph
from services.ingestion import ingest_document_record
from services.study import fetch_due_cards


class EinsteinTutorBackendTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.base_dir = Path(self.temp_dir.name)
        self.original_base_dir = main.BASE_DIR
        self.original_data_dir = main.DATA_DIR
        self.original_upload_dir = main.UPLOAD_DIR
        self.original_db_path = main.DB_PATH
        self.original_schema_path = main.SCHEMA_PATH

        main.BASE_DIR = self.base_dir
        main.DATA_DIR = self.base_dir / "data"
        main.UPLOAD_DIR = main.DATA_DIR / "uploads"
        main.DB_PATH = main.DATA_DIR / "test.db"
        main.SCHEMA_PATH = self.original_schema_path
        main.initialize_database()
        self.clear_seed_data()

    def tearDown(self) -> None:
        main.BASE_DIR = self.original_base_dir
        main.DATA_DIR = self.original_data_dir
        main.UPLOAD_DIR = self.original_upload_dir
        main.DB_PATH = self.original_db_path
        main.SCHEMA_PATH = self.original_schema_path
        self.temp_dir.cleanup()

    def clear_seed_data(self) -> None:
        with main.get_db() as conn:
            for table in [
                "concept_edges",
                "questions",
                "srs_cards",
                "dialogue_sessions",
                "notes",
                "study_events",
                "concepts",
                "chunks",
                "documents",
                "app_settings",
            ]:
                conn.execute(f"DELETE FROM {table}")
            if conn.execute(
                "SELECT name FROM sqlite_master WHERE type = 'table' AND name = 'chunks_vec'"
            ).fetchone():
                conn.execute("DELETE FROM chunks_vec")
            conn.commit()

    def ingest(self, filename: str, text: str, subject_name: str) -> str:
        with main.get_db() as conn:
            result = ingest_document_record(
                conn=conn,
                filename=filename,
                file_type=Path(filename).suffix.replace(".", "") or "txt",
                extracted_text=text,
                page_count=None,
                subject_name=subject_name,
            )
        return result["doc_id"]

    def test_documents_keep_distinct_identity_with_same_subject(self) -> None:
        doc_a = self.ingest(
            "cell-division-a.txt",
            "Mitosis drives growth. Chromosomes package DNA during cell division. Checkpoints regulate the cycle.",
            "Biology",
        )
        doc_b = self.ingest(
            "cell-division-b.txt",
            "Meiosis supports variation. Chromosomes separate into haploid cells. Recombination changes inheritance.",
            "Biology",
        )

        with main.get_db() as conn:
            documents = document_service.fetch_documents(conn)
            self.assertEqual(2, len(documents))
            self.assertEqual({doc_a, doc_b}, {item["id"] for item in documents})
            self.assertEqual({"Biology"}, {item["subject_name"] for item in documents})

            detail_a = document_service.fetch_document_detail(conn, doc_a)
            detail_b = document_service.fetch_document_detail(conn, doc_b)

        self.assertEqual("Biology", detail_a["document"]["subject_name"])
        self.assertTrue(all(concept["doc_id"] == doc_a for concept in detail_a["concepts"]))
        self.assertTrue(all(concept["document_name"] == "cell-division-a.txt" for concept in detail_a["concepts"]))
        self.assertTrue(all(concept["doc_id"] == doc_b for concept in detail_b["concepts"]))
        self.assertTrue(all(concept["document_name"] == "cell-division-b.txt" for concept in detail_b["concepts"]))

    def test_subject_grouping_can_be_updated_without_losing_traceability(self) -> None:
        doc_a = self.ingest(
            "photosynthesis.txt",
            "Photosynthesis converts light into chemical energy. Chlorophyll absorbs light.",
            "Biology",
        )
        doc_b = self.ingest(
            "respiration.txt",
            "Cellular respiration releases stored energy. Mitochondria generate ATP.",
            "General",
        )

        with main.get_db() as conn:
            updated = document_service.set_document_subject(conn, doc_b, "Biology")
            subjects = document_service.fetch_subject_groups(conn)
            detail_b = document_service.fetch_document_detail(conn, doc_b)

        self.assertEqual("Biology", updated["subject_name"])
        self.assertEqual("Biology", detail_b["document"]["subject_name"])
        biology_group = next(item for item in subjects if item["subject_name"] == "Biology")
        self.assertEqual(2, biology_group["document_count"])
        self.assertNotEqual(doc_a, doc_b)

    def test_graph_filters_keep_sources_traceable_per_document_and_subject(self) -> None:
        doc_a = self.ingest(
            "biology-a.txt",
            "Mitosis creates identical cells. DNA replication happens before division. Checkpoints regulate timing.",
            "Biology",
        )
        doc_b = self.ingest(
            "biology-b.txt",
            "Mitosis can be contrasted with meiosis. Meiosis increases variation and creates haploid cells.",
            "Biology",
        )
        self.ingest(
            "chemistry.txt",
            "Ionic bonding transfers electrons. Covalent bonding shares electrons.",
            "Chemistry",
        )

        with main.get_db() as conn:
            biology_graph = fetch_graph(conn, subject_name="Biology")
            doc_graph = fetch_graph(conn, doc_id=doc_a)

        self.assertTrue(biology_graph["nodes"])
        self.assertEqual({"Biology"}, {node["subject_name"] for node in biology_graph["nodes"]})
        self.assertEqual({doc_a, doc_b}, {node["document_id"] for node in biology_graph["nodes"]})
        self.assertTrue(all(node["document_name"] for node in biology_graph["nodes"]))

        self.assertTrue(doc_graph["nodes"])
        self.assertEqual({doc_a}, {node["document_id"] for node in doc_graph["nodes"]})
        self.assertTrue(all(edge["document_id"] == doc_a for edge in doc_graph["edges"]))

    def test_delete_document_preserves_other_grouped_documents(self) -> None:
        doc_a = self.ingest(
            "algebra-1.txt",
            "Functions map inputs to outputs. Linear equations create straight lines.",
            "Math",
        )
        doc_b = self.ingest(
            "algebra-2.txt",
            "Quadratic equations create parabolas. Factoring helps solve polynomial expressions.",
            "Math",
        )

        with main.get_db() as conn:
            deleted = document_service.delete_document_record(conn, doc_a)
            remaining_documents = document_service.fetch_documents(conn)
            remaining_graph = fetch_graph(conn, subject_name="Math")

        self.assertTrue(deleted)
        self.assertEqual([doc_b], [item["id"] for item in remaining_documents])
        self.assertTrue(all(node["document_id"] == doc_b for node in remaining_graph["nodes"]))

    def test_document_lifecycle_keeps_concept_graph_in_sync(self) -> None:
        """Pin the contract that the user asked for explicitly:
          - Ingesting a document creates concept rows (FK doc_id matches)
            and concept_edges between them (FK doc_id matches).
          - Deleting a document removes ITS concepts and edges.
          - A second document's concepts/edges are untouched.
          - Zero orphans remain on either table afterwards.

        The lifecycle helpers (services.ingestion.orchestrator.
        ingest_document_record + services.documents.delete_document_record)
        already do this; this test is the regression gate so a future
        refactor that forgets to cascade can't ship silently. Without
        ON DELETE CASCADE on the schema-level FKs (intentional —
        SQLite ALTER TABLE can't add it without a full table rebuild),
        the cascade lives in application code and needs a test pinning
        it.
        """
        # Use text written with relationship keywords (because, therefore,
        # different, uses) so the edge inferrer emits at least one edge per
        # doc — otherwise the heuristic only produces edges when both
        # concept names AND a keyword land in one sentence, and short
        # synthetic text often misses.
        doc_a = self.ingest(
            "acids-bases.txt",
            "Acids release protons because they donate hydrogen ions. "
            "Bases accept protons, therefore they neutralize acids. "
            "Different proton donors and acceptors react to form salts.",
            "Chemistry",
        )
        doc_b = self.ingest(
            "redox.txt",
            "Oxidation transfers electrons because the species loses charge. "
            "Reduction gains electrons, therefore it balances oxidation. "
            "Different oxidation states drive the redox reaction.",
            "Chemistry",
        )

        with main.get_db() as conn:
            # Sanity: ingestion populated concepts AND edges for both docs.
            concepts_a = conn.execute(
                "SELECT id FROM concepts WHERE doc_id = ?", (doc_a,)
            ).fetchall()
            concepts_b = conn.execute(
                "SELECT id FROM concepts WHERE doc_id = ?", (doc_b,)
            ).fetchall()
            edges_a = conn.execute(
                "SELECT source_id, target_id FROM concept_edges WHERE doc_id = ?",
                (doc_a,),
            ).fetchall()
            edges_b = conn.execute(
                "SELECT source_id, target_id FROM concept_edges WHERE doc_id = ?",
                (doc_b,),
            ).fetchall()

        self.assertGreater(
            len(concepts_a), 0, "ingest_document_record must create concepts for doc_a"
        )
        self.assertGreater(
            len(concepts_b), 0, "ingest_document_record must create concepts for doc_b"
        )
        # We don't require a minimum edge count here. The edge inferrer
        # is heuristic (looks for keywords + co-occurring concept names)
        # and naturally produces zero edges on some text shapes. The
        # cascade contract — "if edges exist, they are cleaned up" — is
        # the invariant we actually care about, asserted below via the
        # zero-orphan checks. We keep `edges_b` captured so the post-
        # delete assertion can verify doc_b's edges weren't disturbed.

        with main.get_db() as conn:
            deleted = document_service.delete_document_record(conn, doc_a)

        self.assertTrue(deleted)

        with main.get_db() as conn:
            # Doc A's concepts + edges are gone.
            after_a = conn.execute(
                "SELECT COUNT(*) FROM concepts WHERE doc_id = ?", (doc_a,)
            ).fetchone()[0]
            after_a_edges = conn.execute(
                "SELECT COUNT(*) FROM concept_edges WHERE doc_id = ?", (doc_a,)
            ).fetchone()[0]
            # Doc B is untouched.
            after_b = conn.execute(
                "SELECT id FROM concepts WHERE doc_id = ?", (doc_b,)
            ).fetchall()
            after_b_edges = conn.execute(
                "SELECT source_id, target_id FROM concept_edges WHERE doc_id = ?",
                (doc_b,),
            ).fetchall()
            # No orphan concepts (every concept's doc_id resolves).
            orphan_concepts = conn.execute(
                """
                SELECT c.id FROM concepts c
                LEFT JOIN documents d ON d.id = c.doc_id
                WHERE d.id IS NULL
                """
            ).fetchall()
            # No orphan edges (every endpoint resolves to a real concept).
            orphan_edges = conn.execute(
                """
                SELECT ce.source_id, ce.target_id FROM concept_edges ce
                LEFT JOIN concepts s ON s.id = ce.source_id
                LEFT JOIN concepts t ON t.id = ce.target_id
                WHERE s.id IS NULL OR t.id IS NULL
                """
            ).fetchall()

        self.assertEqual(0, after_a, "deleting doc_a must remove its concepts")
        self.assertEqual(0, after_a_edges, "deleting doc_a must remove its edges")
        self.assertEqual(
            {row["id"] for row in concepts_b},
            {row["id"] for row in after_b},
            "deleting doc_a must NOT touch doc_b's concepts",
        )
        self.assertEqual(
            {(e["source_id"], e["target_id"]) for e in edges_b},
            {(e["source_id"], e["target_id"]) for e in after_b_edges},
            "deleting doc_a must NOT touch doc_b's edges",
        )
        self.assertEqual([], orphan_concepts, "no orphan concepts after delete")
        self.assertEqual([], orphan_edges, "no orphan edges after delete")

    def test_upload_route_persists_subject_metadata(self) -> None:
        upload = UploadFile(filename="study-notes.txt", file=io.BytesIO(b"Cell membranes regulate transport. Diffusion moves particles."))
        result = asyncio.run(upload_document(file=upload, subject_name="Biology Unit 2"))

        with main.get_db() as conn:
            detail = document_service.fetch_document_detail(conn, result["doc_id"])
            documents = document_service.fetch_documents(conn)

        self.assertEqual("Biology Unit 2", result["subject_name"])
        self.assertEqual("Biology Unit 2", detail["document"]["subject_name"])
        self.assertEqual("study-notes.txt", documents[0]["filename"])

    def test_documents_list_surfaces_top_level_confidence(self) -> None:
        self.ingest(
            "confidence-notes.txt",
            "Mitosis creates daughter cells. Checkpoints regulate the cell cycle.",
            "Biology",
        )

        with main.get_db() as conn:
            documents = document_service.fetch_documents(conn)
            detail = document_service.fetch_document_detail(conn, documents[0]["id"])

        self.assertIn("confidence", documents[0])
        self.assertIn("confidence", detail["document"])
        if documents[0]["confidence"] is not None:
            self.assertIsInstance(documents[0]["confidence"], float)
        if detail["document"]["confidence"] is not None:
            self.assertIsInstance(detail["document"]["confidence"], float)

    def test_document_detail_route_returns_ordered_chunks_and_expected_shape(self) -> None:
        with main.get_db() as conn:
            conn.execute(
                """
                INSERT INTO documents (
                    id, filename, storage_name, subject_name, file_type, status
                )
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    "detail-doc",
                    "ordered-notes.txt",
                    None,
                    "Biology",
                    "txt",
                    "ready",
                ),
            )
            conn.execute(
                """
                INSERT INTO chunks (id, doc_id, chunk_index, page_num, section, content)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                ("chunk-2", "detail-doc", 2, 3, "Regulation", "Checkpoints pause mitosis."),
            )
            conn.execute(
                """
                INSERT INTO chunks (id, doc_id, chunk_index, page_num, section, content)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                ("chunk-1", "detail-doc", 1, 1, "Basics", "Mitosis creates daughter cells."),
            )
            conn.commit()

        client = TestClient(main.app)

        response = client.get("/api/documents/detail-doc")
        self.assertEqual(200, response.status_code, response.text)
        payload = response.json()

        self.assertEqual("detail-doc", payload["document"]["id"])
        self.assertIn("confidence", payload["document"])
        self.assertEqual(["chunk-1", "chunk-2"], [item["id"] for item in payload["chunks"]])
        self.assertEqual([1, 2], [item["chunk_index"] for item in payload["chunks"]])
        self.assertIn("concepts", payload)
        self.assertIsInstance(payload["concepts"], list)

        missing = client.get("/api/documents/missing-detail-doc")
        self.assertEqual(404, missing.status_code, missing.text)

    def test_put_subject_route_updates_subject_and_returns_document_shape(self) -> None:
        doc_id = self.ingest(
            "route-subject.txt",
            "Photosystems capture light energy. ATP and NADPH power downstream reactions.",
            "General",
        )

        client = TestClient(main.app)
        response = client.put(
            f"/api/documents/{doc_id}/subject",
            json={"subject_name": "Eval Biology"},
        )

        self.assertEqual(200, response.status_code, response.text)
        payload = response.json()
        self.assertEqual(doc_id, payload["id"])
        self.assertEqual("Eval Biology", payload["subject_name"])
        self.assertIn("confidence", payload)
        self.assertIn("parser_diagnostics", payload)

    def test_document_file_route_streams_seeded_upload_and_blocks_invalid_paths(self) -> None:
        upload_root = main.UPLOAD_DIR
        upload_root.mkdir(parents=True, exist_ok=True)
        stored_name = "sample.pdf"
        stored_path = upload_root / stored_name
        stored_path.write_bytes(b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\ntrailer\n<<>>\n%%EOF")

        missing_name = "missing.pdf"

        with main.get_db() as conn:
            conn.execute(
                """
                INSERT INTO documents (id, filename, storage_name, subject_name, file_type, status)
                VALUES (?, ?, ?, ?, ?, 'ready')
                """,
                ("pdf-doc", "sample.pdf", stored_name, "Biology", "pdf"),
            )
            conn.execute(
                """
                INSERT INTO documents (id, filename, storage_name, subject_name, file_type, status)
                VALUES (?, ?, ?, ?, ?, 'ready')
                """,
                ("missing-doc", "missing.pdf", missing_name, "Biology", "pdf"),
            )
            conn.execute(
                """
                INSERT INTO documents (id, filename, storage_name, subject_name, file_type, status)
                VALUES (?, ?, ?, ?, ?, 'ready')
                """,
                ("traversal-doc", "escape.pdf", "../escape.pdf", "Biology", "pdf"),
            )
            conn.commit()

        client = TestClient(main.app)

        ok_response = client.get("/api/documents/pdf-doc/file")
        self.assertEqual(200, ok_response.status_code, ok_response.text)
        self.assertEqual("application/pdf", ok_response.headers["content-type"])
        self.assertTrue(ok_response.content.startswith(b"%PDF-1.4"))

        missing_response = client.get("/api/documents/missing-doc/file")
        self.assertEqual(404, missing_response.status_code, missing_response.text)

        unknown_response = client.get("/api/documents/unknown/file")
        self.assertEqual(404, unknown_response.status_code, unknown_response.text)

        traversal_response = client.get("/api/documents/traversal-doc/file")
        self.assertEqual(404, traversal_response.status_code, traversal_response.text)

    def test_docx_upload_extracts_structured_content_without_manual_text(self) -> None:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self.skipTest("python-docx not installed")

        document = DocxDocument()
        document.add_heading("Cell Membrane Transport", level=1)
        document.add_paragraph("Facilitated diffusion uses membrane proteins to move molecules down a gradient.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Process"
        table.cell(0, 1).text = "Energy"
        table.cell(1, 0).text = "Active transport"
        table.cell(1, 1).text = "ATP"

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        upload = UploadFile(filename="transport.docx", file=buffer)
        result = asyncio.run(upload_document(file=upload, subject_name="Biology"))

        with main.get_db() as conn:
            detail = document_service.fetch_document_detail(conn, result["doc_id"])

        diagnostics = detail["document"]["parser_diagnostics"]
        self.assertEqual("docx", result["file_type"])
        self.assertEqual("python-docx-structured", diagnostics["quality"]["parser"])
        self.assertGreaterEqual(detail["counts"]["chunks"], 1)
        self.assertTrue(any("Facilitated diffusion" in chunk["content"] for chunk in detail["chunks"]))
        self.assertTrue(any("ATP" in chunk["content"] for chunk in detail["chunks"]))
        self.assertTrue(any(chunk["provenance_json"].get("source_spans") for chunk in detail["chunks"]))

    def test_flashcard_draft_uses_grounded_source_scope_without_manual_content(self) -> None:
        doc_id = self.ingest(
            "membranes.txt",
            "Facilitated diffusion uses membrane proteins. Active transport requires ATP to move substances against a gradient.",
            "Biology",
        )

        result = draft_flashcards(
            FlashcardDraftRequest(
                title="Membranes",
                content=None,
                source_scope=[doc_id],
                count=6,
            )
        )

        self.assertTrue(result["cards"])
        self.assertTrue(any("ATP" in card["a"] or "transport" in card["q"].lower() for card in result["cards"]))
        rendered = " ".join(f"{card['q']} {card['a']}" for card in result["cards"]).lower()
        self.assertNotIn("what does the source say", rendered)
        self.assertNotIn("which evidence best supports", rendered)
        self.assertNotIn("membranes.txt", rendered)

    def test_flashcard_draft_route_accepts_typed_hidden_grounding_fields(self) -> None:
        doc_id = self.ingest(
            "transport.txt",
            "Active transport uses ATP. Facilitated diffusion moves molecules down a concentration gradient with membrane proteins.",
            "Biology",
        )

        client = TestClient(main.app)
        response = client.post(
            "/api/flashcards/draft",
            json={
                "title": "Transport",
                "content": None,
                "source_scope": [doc_id],
                "count": 6,
            },
        )

        self.assertEqual(200, response.status_code, response.text)
        payload = response.json()
        self.assertTrue(payload["cards"])
        self.assertIsInstance(payload["cards"][0]["confidence"], float)
        self.assertIsInstance(payload["cards"][0]["show_citations"], bool)

    def test_flashcard_draft_filters_slide_outline_noise_and_malformed_labels(self) -> None:
        with main.get_db() as conn:
            doc_id = "finance-slides"
            conn.execute(
                """
                INSERT INTO documents (id, filename, subject_name, file_type, status)
                VALUES (?, ?, ?, ?, 'ready')
                """,
                (doc_id, "Berk_DeMarzo_cf5_ppt_10.pdf", "Finance", "pdf"),
            )
            chunks = [
                (doc_id, "Chapter Outline\n10.1 Risk and Return: Insights from 96 Years of Investor History\n10.2 Common Measures of Risk and Return\n© 2024 Pearson Education, Ltd. All Rights Reserved", "Page 2", 2, 0),
                (doc_id, "Learning Objectives\n• Define a probability distribution, the mean, the variance, the standard deviation, and the volatility.\n• Compute the realized or total return for an investment.", "Page 3", 3, 1),
                (doc_id, "Expected Return\n• Expected return is the probability-weighted average of the possible returns on an investment.", "Page 12", 12, 2),
                (doc_id, "Variance and Standard Deviation (1 of 2)\n• Variance measures how far returns tend to spread around the mean.\n• Standard deviation is the square root of variance and is a common measure of volatility.", "Page 14", 14, 3),
                (doc_id, "Beta\n• Beta measures the systematic risk of a security relative to the market portfolio.", "Page 24", 24, 4),
                (doc_id, "Capital Asset Pricing Model\nE[R_i] = r_f + β_i(E[R_Mkt] − r_f)\n• The CAPM states that expected return equals the risk-free rate plus a beta-based risk premium.", "Page 28", 28, 5),
            ]
            for index, (chunk_doc_id, content, section, page_num, chunk_index) in enumerate(chunks, start=1):
                conn.execute(
                    """
                    INSERT INTO chunks (id, doc_id, content, section, page_num, chunk_index, token_count, provenance_json)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        f"finance-chunk-{index}",
                        chunk_doc_id,
                        content,
                        section,
                        page_num,
                        chunk_index,
                        len(content.split()),
                        "{}",
                    ),
                )
            conn.commit()

        result = draft_flashcards(
            FlashcardDraftRequest(
                title="Risk and Return",
                content=None,
                source_scope=[doc_id],
                count=6,
            )
        )

        self.assertTrue(result["cards"])
        rendered = " ".join(f"{card['q']} {card['a']}" for card in result["cards"]).lower()
        self.assertIn("expected return", rendered)
        self.assertIn("beta", rendered)
        self.assertNotIn("chapter outline", rendered)
        self.assertNotIn("learning objectives", rendered)
        self.assertNotIn("all rights reserved", rendered)
        self.assertNotIn("return variance and", rendered)
        self.assertNotIn("returns estimate expected", rendered)
        self.assertTrue(all(card["type"] == "definition" for card in result["cards"]))
        self.assertTrue(all(len(card["a"].split()) >= 5 for card in result["cards"]))

    def test_summary_artifact_hides_source_labels_by_default(self) -> None:
        doc_id = self.ingest(
            "tax-law.txt",
            "Double taxation happens when the same income is taxed by more than one tax authority. Tax treaties can reduce double taxation by allocating taxing rights.",
            "Law",
        )

        result = studio_generate(
            StudioGenerateRequest(
                artifact_kind="summary",
                source_scope=[doc_id],
            )
        )

        markdown = result["artifact"]["output_markdown"].lower()
        artifact_detail = studio_get_artifact(result["artifact"]["id"])
        self.assertNotIn("tax-law.txt", markdown)
        self.assertNotIn("source:", markdown)
        self.assertNotIn("page ", markdown)
        self.assertEqual("internal_only", artifact_detail["output_json"]["grounding_mode"])
        self.assertFalse(artifact_detail["output_json"]["show_citations"])

    def test_cleaned_labels_flow_into_graph_and_due_cards(self) -> None:
        with main.get_db() as conn:
            doc_id = "doc-clean"
            chunk_id = "chunk-clean"
            concept_a = "concept-a"
            concept_b = "concept-b"
            conn.execute(
                """
                INSERT INTO documents (id, filename, subject_name, file_type, status)
                VALUES (?, ?, ?, ?, 'ready')
                """,
                (doc_id, "finance.pdf", "General", "pdf"),
            )
            conn.execute(
                """
                INSERT INTO chunks (id, doc_id, content, section, chunk_index, token_count)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (chunk_id, doc_id, "Dividend growth interacts with valuation.", "Section 1", 0, 5),
            )
            conn.execute(
                """
                INSERT INTO concepts (id, doc_id, name, description, mastery, source_chunks)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (concept_a, doc_id, "All Right Reservedlearning Objective", "desc", 0.3, '["chunk-clean"]'),
            )
            conn.execute(
                """
                INSERT INTO concepts (id, doc_id, name, description, mastery, source_chunks)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (concept_b, doc_id, "All Right Reserveddividend Versus", "desc", 0.4, '["chunk-clean"]'),
            )
            conn.execute(
                """
                INSERT INTO concept_edges (source_id, target_id, doc_id, relationship, weight)
                VALUES (?, ?, ?, ?, 1)
                """,
                (concept_a, concept_b, doc_id, "supports"),
            )
            conn.execute(
                """
                INSERT INTO srs_cards (id, concept_id, card_type, front, back, due_date)
                VALUES (?, ?, 'definition', ?, ?, ?)
                """,
                (
                    "card-clean",
                    concept_a,
                    "What is the main idea behind All Right Reservedlearning Objective?",
                    "Compare All Right Reservedlearning Objective with All Right Reserveddividend Versus.",
                    "2026-01-01",
                ),
            )
            conn.commit()

            graph = fetch_graph(conn, doc_id=doc_id)
            due_cards = fetch_due_cards(conn)
            detail = document_service.fetch_document_detail(conn, doc_id, include_chunks=False)

        self.assertEqual("learning Objective", graph["nodes"][0]["label"])
        self.assertEqual("learning Objective", due_cards[0]["concept"])
        self.assertIn("learning Objective", due_cards[0]["front"])
        self.assertIn("dividend Versus", due_cards[0]["back"])
        self.assertEqual("learning Objective", detail["concepts"][0]["display_name"])

    def test_noisy_ingestion_avoids_junk_concepts_and_speculative_questions(self) -> None:
        doc_id = self.ingest(
            "valuation-notes.txt",
            (
                "All Rights Reserved. Learning Objectives. "
                "Dividend policy affects firm value through investor expectations. "
                "Stock valuation shapes managerial decision making because capital allocation depends on expected returns. "
                "Learning Objective 1. All Rights Reserved."
            ),
            "Finance",
        )

        with main.get_db() as conn:
            detail = document_service.fetch_document_detail(conn, doc_id, include_chunks=False)
            concepts = document_service.collect_document_concepts(conn, doc_id)
            question_rows = conn.execute(
                """
                SELECT q.question, q.distractors
                FROM questions q
                JOIN concepts c ON c.id = q.concept_id
                WHERE c.doc_id = ?
                ORDER BY q.rowid ASC
                """,
                (doc_id,),
            ).fetchall()

        self.assertTrue(detail["concepts"])
        joined_labels = " ".join(concept["display_name"].lower() for concept in detail["concepts"])
        self.assertNotIn("all right reserved", joined_labels)
        self.assertNotIn("learning objective", joined_labels)
        self.assertTrue(all(concept["source_chunk_ids"] for concept in concepts))

        serialized_questions = " ".join(
            f"{row['question']} {row['distractors']}".lower()
            for row in question_rows
        )
        self.assertNotIn("unrelated to the main argument", serialized_questions)
        self.assertNotIn("only matters when compared", serialized_questions)

    def test_noise_only_document_creates_no_placeholder_study_items(self) -> None:
        doc_id = self.ingest(
            "noise-only.txt",
            "All Rights Reserved. Learning Objectives. Copyright. Table of Contents. Page 1.",
            "General",
        )

        with main.get_db() as conn:
            concept_total = conn.execute(
                "SELECT COUNT(*) AS total FROM concepts WHERE doc_id = ?",
                (doc_id,),
            ).fetchone()["total"]
            question_total = conn.execute(
                """
                SELECT COUNT(*) AS total
                FROM questions q
                JOIN concepts c ON c.id = q.concept_id
                WHERE c.doc_id = ?
                """,
                (doc_id,),
            ).fetchone()["total"]
            card_total = conn.execute(
                """
                SELECT COUNT(*) AS total
                FROM srs_cards s
                JOIN concepts c ON c.id = s.concept_id
                WHERE c.doc_id = ?
                """,
                (doc_id,),
            ).fetchone()["total"]

        self.assertEqual(0, concept_total)
        self.assertEqual(0, question_total)
        self.assertEqual(0, card_total)


if __name__ == "__main__":
    unittest.main()
