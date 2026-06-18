"""Generate the two demo .docx files for the Cachet Wedge-2 contract close.

Hard rule: the load-bearing clause text (dates, amounts, exclusivity, governing
law, the verbatim venue quote) is copied CHARACTER-FOR-CHARACTER from the
verified fixtures (demo/contract-msa.md + demo/contract-ai-summary.md) so the
deterministic engine's anchors fire exactly as in tests/test_demo_corpus.py.
Only anchor-free framing (title, parties, signature block) is added for realism.
No new dates or dollar amounts are introduced anywhere.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INK = RGBColor(0x1A, 0x1A, 0x1A)


def _base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _h(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def _body(doc: Document, text: str, justify: bool = True) -> None:
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _section(doc: Document, label: str, body: str) -> None:
    """One paragraph: a bold section label run followed by the clause body in the
    SAME paragraph. This mirrors demo/contract-msa.md (label + body in one block)
    so Docling keeps the section number with the clause text in a single node and
    the contradiction detail can still name the section (Mythos docx-heading-split).
    """
    # Plain (non-bold) single run: a bold lead run makes Docling split the label
    # into a separate heading node, stripping the section number off the clause
    # body. Keeping it one plain run mirrors the .md and preserves attribution.
    p = doc.add_paragraph(f"{label} {body}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def build_contract(path: str) -> None:
    doc = Document()
    _base_styles(doc)
    _title(doc, "MASTER SERVICES AGREEMENT")
    _body(
        doc,
        'This Master Services Agreement (the "Agreement") is entered into by and '
        "between Meridian Software, Inc., a company organized under the laws of the "
        'State of Delaware ("Licensor"), and Falcon Capital Partners ("Licensee"). '
        'Licensor and Licensee are each a "party" and together the "parties".',
    )
    _body(
        doc,
        "The parties, intending to be legally bound, agree as follows:",
        justify=False,
    )

    # --- Load-bearing sections: clause text is VERBATIM from demo/contract-msa.md ---
    _section(doc, "Section 1. Effective Date.", "This Agreement is dated March 11, 2023.")
    _section(
        doc,
        "Section 2. Services.",
        "Licensor shall provide the Software and related support services to "
        "Licensee in accordance with the terms of this Agreement and any "
        "applicable order form.",
    )
    _section(
        doc,
        "Section 3. License Grant.",
        "Licensor hereby grants Licensee a non-exclusive, non-transferable "
        "license to use the Software during the Term.",
    )
    _section(
        doc,
        "Section 8. Limitation of Liability.",
        "The aggregate liability of either party under this Agreement shall not exceed $500,000.",
    )
    _section(
        doc,
        "Section 12. Term.",
        "This Agreement shall continue for a term of two (2) years from the "
        "Effective Date, unless earlier terminated in accordance with Section 13.",
    )
    _section(
        doc,
        "Section 13. Termination.",
        "Either party may terminate this Agreement for material breach upon "
        "written notice if the breach remains uncured following a reasonable "
        "opportunity to cure.",
    )
    _section(
        doc,
        "Section 14. Governing Law.",
        "This Agreement shall be governed by and construed in accordance with the "
        "laws of the State of Delaware; the parties submit to the exclusive "
        "jurisdiction of the courts of New York.",
    )
    _section(
        doc,
        "Section 15. Confidentiality.",
        "Each party shall use commercially reasonable efforts to protect the other "
        "party's Confidential Information.",
    )

    doc.add_paragraph()
    _body(
        doc,
        "IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.",
        justify=False,
    )
    doc.add_paragraph()
    _body(doc, "LICENSOR: Meridian Software, Inc.", justify=False)
    _body(doc, "By: ______________________________", justify=False)
    doc.add_paragraph()
    _body(doc, "LICENSEE: Falcon Capital Partners", justify=False)
    _body(doc, "By: ______________________________", justify=False)

    doc.save(path)
    print("wrote", path)


def build_summary(path: str) -> None:
    doc = Document()
    _base_styles(doc)
    _title(doc, "MEMORANDUM")
    _body(
        doc,
        "AI-Generated Summary of the Master Services Agreement (the document under verification)",
        justify=False,
    )
    doc.add_paragraph()
    # --- Claim sentences VERBATIM from demo/contract-ai-summary.md ---
    claims = [
        "The Agreement was executed on March 11, 2024.",
        "Liability under this agreement is capped at $1,000,000.",
        "The term of the agreement is two (2) years from the Effective Date.",
        "The vendor grants the customer an exclusive license to use the Software.",
        "The agreement is governed by New York law.",
        'The parties submit to "the exclusive jurisdiction of the courts of New York."',
        "The vendor must use best efforts to protect confidential information.",
    ]
    for c in claims:
        _body(doc, c, justify=False)
    doc.save(path)
    print("wrote", path)


if __name__ == "__main__":
    build_contract("demo/MSA-Meridian-Falcon-EXECUTED.docx")
    build_summary("demo/AI-Summary-Memo.docx")
